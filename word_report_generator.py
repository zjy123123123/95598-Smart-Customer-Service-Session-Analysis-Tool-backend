# -*- coding: utf-8 -*-
"""
Word报告生成器 - 使用模板文件
"""

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from io import BytesIO
import base64
import json
import os
import re
import sys


def resource_path(*parts):
    """Return a path that works in source and PyInstaller bundles."""
    base = getattr(sys, '_MEIPASS', os.path.dirname(os.path.abspath(__file__)))
    return os.path.join(base, *parts)


TEMPLATE_PATH = resource_path('report-template-v2.docx')


def to_float(value, default=0.0):
    """安全转换为浮点数"""
    try:
        return float(value)
    except (ValueError, TypeError):
        return default


def generate_churn_reason_text(reasons, total_churn):
    """根据流失原因数据生成描述文本（仅展示Top2）"""
    if not reasons or total_churn == 0:
        return "无流失数据"
    
    # 过滤掉计数为0的原因
    valid_reasons = {k: v for k, v in reasons.items() if v > 0}
    if not valid_reasons:
        return "无流失数据"
    
    # 关键修复: 分母使用该环节流失数(total_churn),而非原因值之和
    # 原因: 同一通电话的B列文本可能匹配多个原因关键词,导致原因值之和 > 实际流失数
    total = total_churn
    
    # 按数量降序排序，只取前2
    sorted_reasons = sorted(valid_reasons.items(), key=lambda x: x[1], reverse=True)[:2]
    
    # 生成描述
    nums = ['一', '二']
    parts = []
    for i, (reason, count) in enumerate(sorted_reasons):
        pct = (count / total) * 100
        parts.append(f"{nums[i]}是{reason}（{pct:.1f}%，{count}通）")
    
    return "，".join(parts) + "。"


def _insert_identity_type_table(doc, data):
    """在身份验证段落后插入五类信息统计表格，风格与客户流失率整体情况表完全一致"""
    identity_stats = data.get('identityTypeStats', [])
    if not identity_stats or not isinstance(identity_stats, list) or len(identity_stats) == 0:
        return
    identity_para = None
    for para in doc.paragraphs:
        if '身份验证环节共' in para.text and '流失原因主要包括' in para.text:
            identity_para = para
            break
    if identity_para is None:
        return
    identity_elem = identity_para._element
    
    # === 创建引导句 ===
    guidance_p = OxmlElement('w:p')
    old_pPr = identity_elem.find(qn('w:pPr'))
    if old_pPr is not None:
        import copy
        guidance_pPr = copy.deepcopy(old_pPr)
        guidance_p.append(guidance_pPr)
    guidance_r = OxmlElement('w:r')
    old_rPr = identity_elem.find('.//' + qn('w:rPr'))
    if old_rPr is not None:
        guidance_rPr = copy.deepcopy(old_rPr)
        guidance_r.append(guidance_rPr)
    guidance_t = OxmlElement('w:t')
    guidance_t.text = '各类信息提交及验证通过情况详见下表：'
    guidance_t.set(qn('xml:space'), 'preserve')
    guidance_r.append(guidance_t)
    guidance_p.append(guidance_r)
    
    # === 创建表格 XML 元素（与原始 Table 1 样式完全一致） ===
    def _make_cell(text, shading_fill, font_color, is_bold):
        tc = OxmlElement('w:tc')
        tcPr = OxmlElement('w:tcPr')
        tcW = OxmlElement('w:tcW')
        tcW.set(qn('w:w'), '1701')
        tcW.set(qn('w:type'), 'dxa')
        tcPr.append(tcW)
        shd = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear')
        shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), shading_fill)
        tcPr.append(shd)
        va = OxmlElement('w:vAlign')
        va.set(qn('w:val'), 'center')
        tcPr.append(va)
        tc.append(tcPr)
        p = OxmlElement('w:p')
        pPr = OxmlElement('w:pPr')
        spacing = OxmlElement('w:spacing')
        spacing.set(qn('w:before'), '20')
        spacing.set(qn('w:after'), '20')
        spacing.set(qn('w:line'), '240')
        spacing.set(qn('w:lineRule'), 'auto')
        pPr.append(spacing)
        jc = OxmlElement('w:jc')
        jc.set(qn('w:val'), 'center')
        pPr.append(jc)
        p.append(pPr)
        r = OxmlElement('w:r')
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), '黑体')
        rFonts.set(qn('w:hAnsi'), '黑体')
        rFonts.set(qn('w:eastAsia'), '黑体')
        rPr.append(rFonts)
        if is_bold:
            b = OxmlElement('w:b')
            rPr.append(b)
        color = OxmlElement('w:color')
        color.set(qn('w:val'), font_color)
        rPr.append(color)
        sz = OxmlElement('w:sz')
        sz.set(qn('w:val'), '21')
        rPr.append(sz)
        r.append(rPr)
        t = OxmlElement('w:t')
        t.text = str(text)
        t.set(qn('xml:space'), 'preserve')
        r.append(t)
        p.append(r)
        tc.append(p)
        return tc
    
    def _make_tr():
        tr = OxmlElement('w:tr')
        trPr = OxmlElement('w:trPr')
        tr_jc = OxmlElement('w:jc')
        tr_jc.set(qn('w:val'), 'center')
        trPr.append(tr_jc)
        tr.append(trPr)
        return tr
    
    # 构建表格 XML 结构
    table_elem = OxmlElement('w:tbl')
    tblPr = OxmlElement('w:tblPr')
    ts = OxmlElement('w:tblStyle')
    ts.set(qn('w:val'), '7')
    tblPr.append(ts)
    tw = OxmlElement('w:tblW')
    tw.set(qn('w:w'), '0')
    tw.set(qn('w:type'), 'auto')
    tblPr.append(tw)
    jc_e = OxmlElement('w:jc')
    jc_e.set(qn('w:val'), 'center')
    tblPr.append(jc_e)
    tl = OxmlElement('w:tblLayout')
    tl.set(qn('w:type'), 'fixed')
    tblPr.append(tl)
    tcm = OxmlElement('w:tblCellMar')
    for side, val in [('top', '0'), ('left', '108'), ('bottom', '0'), ('right', '108')]:
        el = OxmlElement('w:' + side)
        el.set(qn('w:w'), val)
        el.set(qn('w:type'), 'dxa')
        tcm.append(el)
    tblPr.append(tcm)
    table_elem.append(tblPr)
    
    # 添加网格列定义
    tblGrid = OxmlElement('w:tblGrid')
    for _ in range(5):
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), '1701')
        tblGrid.append(gc)
    table_elem.append(tblGrid)
    
    # 表头行（深绿底 #006B54，白色字，加粗）
    headers = ['序号', '信息类型', '提交条数', '验证通过', '通过率']
    tr_header = _make_tr()
    for h in headers:
        tr_header.append(_make_cell(h, '006B54', 'FFFFFF', True))
    table_elem.append(tr_header)
    
    # 数据行（浅绿底 #A5D6A7，黑色字，不加粗）
    for ri, stats in enumerate(identity_stats):
        vals = [str(ri + 1), str(stats.get('name', '')), str(stats.get('total', '')), str(stats.get('success', '')), str(stats.get('rate', ''))]
        tr_data = _make_tr()
        for v in vals:
            tr_data.append(_make_cell(v, 'auto', '000000', False))
        table_elem.append(tr_data)
    
    # 插入到身份验证段落后
    identity_elem.addnext(guidance_p)
    guidance_p.addnext(table_elem)


def replace_all_placeholders(doc, data):
    # 调试：打印接收到的关键数据
    print(f'[DEBUG] 接收到的 resultChurnRate: {data.get("resultChurnRate")}')
    print(f'[DEBUG] 接收到的 identityChurnRate: {data.get("identityChurnRate")}')
    print(f'[DEBUG] 接收到的 completeRate: {data.get("completeRate")}')
    print(f'[DEBUG] resultChurnRate 类型: {type(data.get("resultChurnRate"))}')
    """替换文档中的所有占位符（包括表格）"""
    
    # 生成详细流失原因描述
    appeal_reasons = data.get('appealReasons', {})
    relation_reasons = data.get('relationReasons', {})
    identity_reasons = data.get('identityReasons', {})
    result_reasons = data.get('resultReasons', {})
    
    # 各环节流失原因独立展示，不跨环节借用数据
    appeal_churn = to_float(data.get('appealChurn'))
    relation_churn = to_float(data.get('relationChurn'))
    identity_churn = to_float(data.get('identityChurn'))
    result_churn = to_float(data.get('resultChurn'))
    
    # 找出流失率最高的环节
    churn_rates = {
        '诉求接入': to_float(data.get('appealChurnRate')),
        '关系库选择': to_float(data.get('relationChurnRate')),
        '身份验证': to_float(data.get('identityChurnRate')),
        '结果播报': to_float(data.get('resultChurnRate'))
    }
    highest_loss_stage = max(churn_rates, key=churn_rates.get)
    highest_loss_rate = churn_rates[highest_loss_stage]
    
    # 构建替换映射
    replacements = {
        '{{reportDate}}': str(data.get('reportDate', '')),
        '{{sceneName}}': str(data.get('sceneName', '停电信息查询')),
        '{{sessionCount}}': str(data.get('sessionCount', '')),
        '{{lossRate}}': f"{to_float(data.get('lossRate')):.1f}",
        '{{completeRate}}': f"{to_float(data.get('completeRate')):.1f}",
        '{{transferRate}}': f"{to_float(data.get('transferRate')):.1f}",
        '{{appealCount}}': str(data.get('appealCount', '')),
        '{{appealChurn}}': str(data.get('appealChurn', '')),
        '{{appealChurnRate}}': f"{to_float(data.get('appealChurnRate')):.1f}",
        '{{relationCount}}': str(data.get('relationCount', '')),
        '{{relationChurn}}': str(data.get('relationChurn', '')),
        '{{relationChurnRate}}': f"{to_float(data.get('relationChurnRate')):.1f}",
        '{{identityCount}}': str(data.get('identityCount', '')),
        '{{identityChurn}}': str(data.get('identityChurn', '')),
        '{{identityChurnRate}}': f"{to_float(data.get('identityChurnRate')):.1f}",
        '{{resultCount}}': str(data.get('resultCount', '')),
        '{{resultChurn}}': str(data.get('resultChurn', '')),
        '{{resultChurnRate}}': f"{to_float(data.get('resultChurnRate')):.1f}",
        '{{highestLossStage}}': highest_loss_stage,
        '{{highestLossRate}}': f"{highest_loss_rate:.1f}",
        # 流转数据
        '{{flowToIdentity}}': str(data.get('flowToIdentity', 0)),
        '{{flowToBroadcast}}': str(data.get('flowToBroadcast', 0)),
        # 详细流失原因描述
        '{{relationReasons}}': generate_churn_reason_text(relation_reasons, relation_churn),
        '{{identityReasons}}': generate_churn_reason_text(identity_reasons, identity_churn),
        '{{appealReasons}}': generate_churn_reason_text(appeal_reasons, appeal_churn),
        '{{resultReasons}}': generate_churn_reason_text(result_reasons, result_churn),
        # 环节评估（由AI生成）
        '{{appealEval}}': str(data.get('appealEval', '')),
        '{{relationEval}}': str(data.get('relationEval', '')),
        '{{identityEval}}': str(data.get('identityEval', '')),
        '{{resultEval}}': str(data.get('resultEval', '')),
    }
    
    # 替换段落中的占位符（使用段落级替换，处理跨run的占位符）
    for para in doc.paragraphs:
        para_text = para.text
        has_replacement = False
        is_eval_para = False
        for key, value in replacements.items():
            if key in para_text:
                para_text = para_text.replace(key, value)
                has_replacement = True
                if 'Eval' in key:
                    is_eval_para = True
        
        # 特殊处理：关系库选择环节段落 - 严格按标准模板格式生成
        if '关系库选择环节共' in para_text and '流失原因主要包括' in para_text:
            # 提取所有需要的数据
            session_count = data.get('relationCount', '')
            push1 = data.get('push1', 0)
            push2 = data.get('push2', 0)
            push3 = data.get('push3', 0)
            push4plus = data.get('push4plus', 0)
            flow_to_identity = data.get('flowToIdentity', 0)
            flow_to_broadcast = data.get('flowToBroadcast', 0)
            relation_churn = data.get('relationChurn', '')
            relation_churn_rate = data.get('relationChurnRate', '')
            
            # 从 relation_reasons 中提取各原因数量和占比（仅Top2）
            reasons_text_parts = []
            if relation_reasons:
                # 过滤计数为0的原因
                valid_reasons = {k: v for k, v in relation_reasons.items() if v > 0}
                if valid_reasons:
                    # 分母使用该环节流失数
                    total_reason_count = to_float(relation_churn)
                    if total_reason_count == 0:
                        total_reason_count = sum(valid_reasons.values())
                    sorted_reasons = sorted(valid_reasons.items(), key=lambda x: x[1], reverse=True)[:2]
                    nums = ['一', '二']
                    for idx2, (reason, count) in enumerate(sorted_reasons):
                        pct = (count / total_reason_count) * 100
                        reasons_text_parts.append(f"{nums[idx2]}是{reason}（{pct:.1f}%，{count}通）")
            if reasons_text_parts:
                reasons_text = '，'.join(reasons_text_parts)
            else:
                reasons_text = '无流失数据'
            
            # 按标准模板拼接整段话
            new_para = f"关系库选择环节共{session_count}条会话，推送1个户号{push1}条、推送2个户号{push2}条、推送3个户号{push3}条；其中流向身份验证环节{flow_to_identity}条，流向结果播报环节{flow_to_broadcast}条，流失{relation_churn}条，流失率为{relation_churn_rate}%。流失原因主要包括：{reasons_text}。"
            
            para_text = new_para
        
        if has_replacement and para.runs:
            if is_eval_para or para_text.startswith('环节评估'):
                # 环节评估：「环节评估：/环节评估:」加粗，后面描述文字不加粗，16pt，仿宋_GB2312
                print(f'[DEBUG-Eval] 处理环节评估段落, para_text长度={len(para_text)}, runs数量={len(para.runs)}')
                print(f'[DEBUG-Eval] para_text开头50字: {para_text[:50]!r}')
                
                # 确保至少有两个 run
                while len(para.runs) < 2:
                    para.add_run()
                
                # 找第一个加粗 run 和第一个不加粗 run
                bold_run = None
                normal_run = None
                for r in para.runs:
                    if bold_run is None and r.bold:
                        bold_run = r
                    if normal_run is None and not r.bold:
                        normal_run = r
                
                # 兜底：如果找不到，直接用前两个 run
                if bold_run is None:
                    bold_run = para.runs[0]
                    print(f'[DEBUG-Eval] 未找到bold run，使用run0')
                if normal_run is None:
                    for r in para.runs:
                        if r is not bold_run:
                            normal_run = r
                            break
                    if normal_run is None:
                        normal_run = para.add_run()
                    print(f'[DEBUG-Eval] 未找到normal run，创建新run')
                
                # 用底层 XML 元素 ID 做身份比较（避免 Run 包装对象每次重新创建导致 is 比较失效）
                bold_elem_id = id(bold_run._element)
                normal_elem_id = id(normal_run._element)
                print(f'[DEBUG-Eval] bold_elem_id={bold_elem_id}, normal_elem_id={normal_elem_id}')
                
                # 拆分文本（兼容半角和全角冒号）
                if '环节评估：' in para_text:
                    prefix = '环节评估：'
                    content = para_text.split('环节评估：', 1)[1]
                elif '环节评估:' in para_text:
                    prefix = '环节评估:'
                    content = para_text.split('环节评估:', 1)[1]
                else:
                    prefix = para_text
                    content = ''
                
                print(f'[DEBUG-Eval] prefix={prefix!r}, content长度={len(content)}')
                
                # 设置加粗 run（前缀）
                bold_run.text = prefix
                bold_run.bold = True
                bold_run.font.size = Pt(16)
                bold_run.font.name = '仿宋_GB2312'
                rPr = bold_run._element.get_or_add_rPr()
                rFonts = rPr.find(qn('w:rFonts'))
                if rFonts is None:
                    rFonts = rPr.makeelement(qn('w:rFonts'), {})
                    rPr.insert(0, rFonts)
                rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                rFonts.set(qn('w:ascii'), '仿宋_GB2312')
                rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
                rFonts.set(qn('w:cs'), '仿宋_GB2312')
                
                # 设置不加粗 run（描述正文）
                normal_run.text = content
                normal_run.bold = False
                normal_run.font.size = Pt(16)
                normal_run.font.name = '仿宋_GB2312'
                rPr2 = normal_run._element.get_or_add_rPr()
                rFonts2 = rPr2.find(qn('w:rFonts'))
                if rFonts2 is None:
                    rFonts2 = rPr2.makeelement(qn('w:rFonts'), {})
                    rPr2.insert(0, rFonts2)
                rFonts2.set(qn('w:eastAsia'), '仿宋_GB2312')
                rFonts2.set(qn('w:ascii'), '仿宋_GB2312')
                rFonts2.set(qn('w:hAnsi'), '仿宋_GB2312')
                rFonts2.set(qn('w:cs'), '仿宋_GB2312')
                
                # 清空其余 run（用 XML 元素 ID 比较，不用 Run 对象 is 比较）
                for r in para.runs:
                    elem_id = id(r._element)
                    if elem_id != bold_elem_id and elem_id != normal_elem_id:
                        r.text = ''
                
                # 验证最终结果
                final_text = ''.join(r.text for r in para.runs)
                print(f'[DEBUG-Eval] 最终段落文本长度={len(final_text)}, 开头50字: {final_text[:50]!r}')
            else:
                # 非环节评估段落：检查是否有bold分段（如小标题加粗、正文不加粗）
                bold_runs = [r for r in para.runs if r.bold == True]
                normal_runs = [r for r in para.runs if r.bold == False]
                
                if bold_runs and normal_runs and '。' in para_text:
                    # 有bold分段：按第一个'。'分割文本
                    first_period = para_text.index('。') + 1
                    subtitle = para_text[:first_period]
                    content = para_text[first_period:]
                    
                    # 将小标题放入第一个bold run
                    if bold_runs:
                        bold_runs[0].text = subtitle
                        for run in bold_runs[1:]:
                            run.text = ''
                    
                    # 将正文放入第一个normal run
                    if normal_runs:
                        normal_runs[0].text = content
                        for run in normal_runs[1:]:
                            run.text = ''
                else:
                    # 没有bold分段：使用原逻辑，全部放入第一个run
                    para.runs[0].text = para_text
                    for run in para.runs[1:]:
                        run.text = ''
    
    # 插入身份验证环节五类信息统计表格
    _insert_identity_type_table(doc, data)
    
    # 替换表格中的占位符
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para_text = para.text
                    has_replacement = False
                    for key, value in replacements.items():
                        if key in para_text:
                            para_text = para_text.replace(key, value)
                            has_replacement = True
                    
                    if has_replacement and para.runs:
                        para.runs[0].text = para_text
                        for run in para.runs[1:]:
                            run.text = ''
    
    # 更新表格数据（第3行，索引2）
    if len(doc.tables) > 0:
        table = doc.tables[0]
        if len(table.rows) >= 3:
            row = table.rows[2]
            # 第3行: 场景 | 转人工率 | 诉求接入 | 关系库选择 | 身份验证 | 结果播报 | 流程完成率
            if len(row.cells) >= 7:
                # 更新数据
                transfer_rate_val = f"{to_float(data.get('transferRate', 0.0)):.1f}%"
                appeal_rate = f"{to_float(data.get('appealChurnRate')):.1f}%"
                relation_rate = f"{to_float(data.get('relationChurnRate')):.1f}%"
                identity_rate = f"{to_float(data.get('identityChurnRate')):.1f}%"
                result_rate_raw = data.get('resultChurnRate')
                print(f'[DEBUG] 表格更新 - resultChurnRate 原始值: {result_rate_raw}, 类型: {type(result_rate_raw)}')
                result_rate = f"{to_float(data.get('resultChurnRate')):.1f}%"
                print(f'[DEBUG] 表格更新 - result_rate 格式化后: {result_rate}')
                complete_rate = f"{to_float(data.get('completeRate')):.1f}%"
                
                def update_cell_text(cell, new_text, col_name):
                    """更新单元格文本"""
                    print(f"  {col_name}: 更新为 '{new_text}'")
                    # 直接替换整个单元格的文本（保留格式）
                    for para in cell.paragraphs:
                        if para.runs:
                            # 在第一个run中设置新值
                            para.runs[0].text = new_text
                            # 清空其他runs
                            for run in para.runs[1:]:
                                run.text = ''
                
                # 修复：不依赖单元格索引（合并单元格会导致索引偏移）
                # 遍历第3行所有单元格，根据单元格内的文本内容判断应该更新哪个值
                print(f'[DEBUG] 第3行单元格数量: {len(row.cells)}')
                for ci, cell in enumerate(row.cells):
                    cell_text = ''
                    for para in cell.paragraphs:
                        for run in para.runs:
                            cell_text += run.text
                    print(f'[DEBUG] 单元格{ci} 原始内容: "{cell_text.strip()}"')
                
                # 直接按索引更新（模板第3行有7个独立单元格）
                print(f'[DEBUG] 开始更新表格数据...')
                update_cell_text(row.cells[1], transfer_rate_val, "转人工率")
                update_cell_text(row.cells[2], appeal_rate, "诉求接入")
                update_cell_text(row.cells[3], relation_rate, "关系库选择")
                update_cell_text(row.cells[4], identity_rate, "身份验证")
                update_cell_text(row.cells[5], result_rate, "结果播报")
                update_cell_text(row.cells[6], complete_rate, "流程完成率")

                # === 设置客户流失率表格样式（复用身份验证表格的表头颜色） ===
                def _set_cell_shading(cell, fill_color):
                    """设置单元格背景填充色 — 先删除原有shd再重建，确保Word一定生效"""
                    tc = cell._element
                    tcPr = tc.find(qn('w:tcPr'))
                    if tcPr is None:
                        tcPr = OxmlElement('w:tcPr')
                        tc.insert(0, tcPr)
                    # 先删除已有的 shd 元素（清除模板自带样式）
                    old_shd = tcPr.find(qn('w:shd'))
                    if old_shd is not None:
                        tcPr.remove(old_shd)
                    # 重建 shd 元素
                    shd = OxmlElement('w:shd')
                    shd.set(qn('w:val'), 'clear')
                    shd.set(qn('w:color'), 'auto')
                    shd.set(qn('w:fill'), fill_color)
                    tcPr.append(shd)

                def _set_cell_font(cell, color, bold):
                    """设置单元格字体颜色和加粗"""
                    for para in cell.paragraphs:
                        for run in para.runs:
                            rPr = run._element.find(qn('w:rPr'))
                            if rPr is None:
                                from docx.oxml.ns import qn as qn2
                                rPr = OxmlElement('w:rPr')
                                run._element.insert(0, rPr)
                            # 颜色
                            color_elem = rPr.find(qn('w:color'))
                            if color_elem is None:
                                color_elem = OxmlElement('w:color')
                                rPr.append(color_elem)
                            color_elem.set(qn('w:val'), color)
                            # 加粗
                            b_elem = rPr.find(qn('w:b'))
                            if bold and b_elem is None:
                                b_elem = OxmlElement('w:b')
                                rPr.append(b_elem)
                            elif not bold and b_elem is not None:
                                rPr.remove(b_elem)
                            # 字体
                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is None:
                                rFonts = OxmlElement('w:rFonts')
                                rPr.append(rFonts)
                            rFonts.set(qn('w:ascii'), '黑体')
                            rFonts.set(qn('w:hAnsi'), '黑体')
                            rFonts.set(qn('w:eastAsia'), '黑体')
                            sz = rPr.find(qn('w:sz'))
                            if sz is None:
                                sz = OxmlElement('w:sz')
                                rPr.append(sz)
                            sz.set(qn('w:val'), '21')

                # 表头行（第0-1行）：深绿底 #006B54，白字，加粗
                print(f'[DEBUG] 表格总行数: {len(table.rows)}, 开始设置表头样式')
                for header_row in table.rows[0:2]:
                    # 先设置单元格填充（XML）— 不删除原有shd，只修改fill值，保留模板结构
                    for tc in header_row._tr.findall(qn('w:tc')):
                        tcPr = tc.find(qn('w:tcPr'))
                        if tcPr is None:
                            tcPr = OxmlElement('w:tcPr')
                            tc.insert(0, tcPr)
                        # 找到或创建 shd 元素
                        shd = tcPr.find(qn('w:shd'))
                        if shd is None:
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:val'), 'clear')
                            shd.set(qn('w:color'), 'auto')
                            tcPr.append(shd)
                        # 只修改fill值，保留val和color的原始设置
                        shd.set(qn('w:fill'), '006B54')
                    # 再用 python-docx API 设置字体（白字加粗）
                    for cell in header_row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                                run.font.bold = True
                                run.font.name = '黑体'
                                run.font.size = Pt(10.5)
                    print(f'[DEBUG] 表头行已处理')

                # 数据行（第2行起）：无填充，黑字，不加粗
                for data_row in table.rows[2:]:
                    for tc in data_row._tr.findall(qn('w:tc')):
                        tcPr = tc.find(qn('w:tcPr'))
                        if tcPr is None:
                            tcPr = OxmlElement('w:tcPr')
                            tc.insert(0, tcPr)
                        # 找到或创建 shd 元素
                        shd = tcPr.find(qn('w:shd'))
                        if shd is None:
                            shd = OxmlElement('w:shd')
                            shd.set(qn('w:val'), 'clear')
                            shd.set(qn('w:color'), 'auto')
                            tcPr.append(shd)
                        shd.set(qn('w:fill'), 'auto')
                    for cell in data_row.cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                                run.font.bold = False

    # === 清理环节评估段落后方的多余空段落 ===
    # 规则：仅删除紧跟"环节评估:"段落后面的纯空段落（0个run或纯空白文本）
    # 保护：有文字/表格/标题的段落一律保留
    _clean_eval_blank_paras(doc)


def _clean_eval_blank_paras(doc):
    """清理四个环节评估段落后方的纯空段落，消除末尾多余空行"""
    # 从后往前遍历，避免删除后索引偏移
    eval_paras = []
    for i, para in enumerate(doc.paragraphs):
        if '环节评估：' in para.text or '环节评估:' in para.text:
            eval_paras.append(i)
    
    deleted_count = 0
    for eval_idx in reversed(eval_paras):
        next_idx = eval_idx + 1
        if next_idx >= len(doc.paragraphs):
            continue
        next_para = doc.paragraphs[next_idx]
        # 判断是否为纯空段落：无文本 + 无run或run全空
        is_blank = not next_para.text.strip() and (
            len(next_para.runs) == 0 or
            all(not run.text.strip() for run in next_para.runs)
        )
        if is_blank:
            # 从XML树中移除该段落
            next_para._element.getparent().remove(next_para._element)
            deleted_count += 1
    
    if deleted_count > 0:
        print(f"[Report] 清理了 {deleted_count} 个环节评估后的空段落")


def add_funnel_image(doc, funnel_base64):
    """替换模板中的漏斗图（删除旧图片，插入新图片）"""
    if not funnel_base64:
        return
    
    try:
        # 解码base64图片
        if ',' in funnel_base64:
            funnel_base64 = funnel_base64.split(',')[1]
        image_data = base64.b64decode(funnel_base64)
        
        # 找到包含图片的段落（在"2.各环节具体流失原因分析"之后）
        found_section = False
        for i, para in enumerate(doc.paragraphs):
            if '2.各环节具体流失原因分析' in para.text:
                found_section = True
                continue
            
            if found_section:
                # 检查这个段落是否包含图片
                has_image = False
                for run in para.runs:
                    drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                    if drawings:
                        has_image = True
                        break
                
                if has_image:
                    # 删除旧图片
                    for run in para.runs:
                        drawings = run._element.findall('.//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}drawing')
                        for drawing in drawings:
                            drawing.getparent().remove(drawing)
                    
                    # 插入新图片
                    run = para.runs[0] if para.runs else para.add_run()
                    image_stream = BytesIO(image_data)
                    run.add_picture(image_stream, width=Inches(6.0))
                    break
    except Exception as e:
        print(f"[Report] 漏斗图替换失败: {e}")


def set_heading_outline_level(doc):
    """设置一级标题段落的outline level，使其出现在目录中"""
    for para in doc.paragraphs:
        text = para.text.strip()
        # 匹配"一、停电信息查询场景客户流失情况分析"或"停电信息查询场景客户流失情况分析"
        if ('客户流失情况分析' in text or '场景优化建议' in text) and len(text) < 30:
            # 设置outline level为0（一级标题）
            pPr = para._element.get_or_add_pPr()
            outlineLvl = pPr.makeelement(qn('w:outlineLvl'), {qn('w:val'): '0'})
            # 移除已有的outlineLvl
            existing = pPr.findall(qn('w:outlineLvl'))
            for elem in existing:
                pPr.remove(elem)
            pPr.append(outlineLvl)
            print(f"[Report] 设置一级标题: {text}")


def replace_ai_suggestions(doc, ai_text):
    """替换第三章的AI优化建议内容（只替换文本，保持模板原有格式）"""
    if not ai_text:
        print("[Report] AI建议为空，保留模板内容")
        return
    
    print(f"[Report] 开始替换AI建议，文本长度: {len(ai_text)}")
    
    # 解析AI建议文本，按维度分割
    dimensions = []
    current_dim = None
    current_content = []
    
    lines = ai_text.split('\n')
    for line in lines:
        line = line.strip()
        if not line:
            continue
        
        # 检测维度标题
        if line.startswith('【') and line.endswith('】'):
            if current_dim:
                dimensions.append((current_dim, current_content))
            current_dim = line[1:-1]
            current_content = []
        elif line.startswith('###'):
            if current_dim:
                dimensions.append((current_dim, current_content))
            current_dim = line.replace('###', '').strip()
            current_content = []
        elif line.startswith('**') and '**' in line[2:]:
            if current_dim:
                dimensions.append((current_dim, current_content))
            current_dim = line.replace('**', '').strip()
            current_content = []
        else:
            current_content.append(line)
        
        # 统一处理：去掉dim_name中的数字编号前缀（如"3.知识度与话术策略优化"→"知识度与话术策略优化"）
        if current_dim and current_dim[0].isdigit() and '.' in current_dim[:3]:
            dot_pos = current_dim.index('.')
            current_dim = current_dim[dot_pos + 1:].strip()
    
    if current_dim:
        dimensions.append((current_dim, current_content))
    
    print(f"[Report] 解析出 {len(dimensions)} 个维度")
    
    # 找到第三章
    suggestion_idx = -1
    for i, para in enumerate(doc.paragraphs):
        if '三、场景优化建议' in para.text:
            suggestion_idx = i
            break
    
    if suggestion_idx == -1:
        print("[Report] 未找到第三章标题")
        return
    
    # 收集第三章的内容段落
    # 修复：扩大收集范围，不要遇到"一、""二、"就停止，因为这些可能是AI建议内容的一部分
    # 只检查真正的大章节标题（如"一、停电信息查询场景客户流失情况分析"）
    content_paras = []
    for i in range(suggestion_idx + 1, len(doc.paragraphs)):
        para = doc.paragraphs[i]
        text = para.text.strip()
        # 只检查真正的大章节标题（如"一、停电信息查询场景客户流失情况分析"）
        # 不检查AI建议内部的"1.""2."编号
        if any(marker in text for marker in ['一、停电', '二、', '四、', '五、']) and len(text) < 30:
            break
        if text:
            content_paras.append((i, para))
    
    print(f"[Report] 模板中有 {len(content_paras)} 个内容段落")
    
    # 替换内容段落 - 复用模板的 run，只改文本
    for idx, (para_idx, para) in enumerate(content_paras):
        if idx < len(dimensions):
            dim_name, content_lines = dimensions[idx]
            
            # 复用模板的第一个 run（保留原有格式）
            if para.runs:
                first_run = para.runs[0]
                # 只修改文本，不改变格式
                first_run.text = f"{idx + 1}.{dim_name}。"
                # 保持模板的字体格式，只设置加粗
                first_run.bold = True
                # 确保字体和字号与模板一致（三号 = 16pt）
                first_run.font.size = Pt(16)
                first_run.font.name = '仿宋_GB2312'
                # 设置中文字体
                if first_run.font.element is not None:
                    rPr = first_run.font.element.get_or_add_rPr()
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is not None:
                        rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
            
            # 如果有后续 run，复用它们来存放正文内容
            if content_lines:
                if len(para.runs) > 1:
                    # 复用第二个 run
                    content_run = para.runs[1]
                    content_run.text = ''.join(content_lines)
                    content_run.bold = False
                    content_run.font.size = Pt(16)
                    content_run.font.name = '仿宋_GB2312'
                    if content_run.font.element is not None:
                        rPr = content_run.font.element.get_or_add_rPr()
                        rFonts = rPr.find(qn('w:rFonts'))
                        if rFonts is not None:
                            rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                else:
                    # 没有后续 run，创建一个新的（继承第一个 run 的格式）
                    if para.runs:
                        ref_run = para.runs[0]
                        content_run = para.add_run(''.join(content_lines))
                        # 复制参考 run 的格式
                        content_run.font.name = ref_run.font.name
                        content_run.font.size = ref_run.font.size
                        content_run.font.bold = False
                        if content_run.font.element is not None:
                            rPr = content_run.font.element.get_or_add_rPr()
                            rFonts = rPr.find(qn('w:rFonts'))
                            if rFonts is not None:
                                rFonts.set(qn('w:eastAsia'), ref_run.font.name or '仿宋_GB2312')
            
            # 清除多余的 run（如果有的话）
            for run in para.runs[2:]:
                run.text = ''
            
            print(f"[Report] 替换段落 {para_idx}: 标题={idx + 1}.{dim_name}")
        else:
            # 没有对应的AI维度，清空段落
            if para.runs:
                for run in para.runs:
                    run.text = ''
            print(f"[Report] 清空段落 {para_idx}")
    
    # 修复：如果AI维度多于模板段落，需要添加新段落
    if len(dimensions) > len(content_paras):
        print(f"[Report] 警告：AI建议有{len(dimensions)}个维度，但模板只有{len(content_paras)}个段落")
        # 在最后一个段落后面添加新段落，复制模板样式
        if content_paras:
            last_para_idx = content_paras[-1][0]
            last_para = doc.paragraphs[last_para_idx]
            # 获取最后一个段落在文档中的位置索引
            body = doc.element.body
            last_element = last_para._element
            for idx in range(len(content_paras), len(dimensions)):
                dim_name, content_lines = dimensions[idx]
                # 用XML方式在last_element后面插入新段落
                new_p = OxmlElement('w:p')
                # 完整克隆模板段落的pPr（包含行距、缩进等所有格式）
                last_pPr = last_para._element.find(qn('w:pPr'))
                if last_pPr is not None:
                    import copy
                    new_pPr = copy.deepcopy(last_pPr)
                    new_p.append(new_pPr)
                
                body.append(new_p)
                new_para = Paragraph(new_p, doc)
                
                # 按"。"分割，小标题加粗，正文不加粗
                content_text = ''.join(content_lines)
                full_text = f"{idx + 1}.{dim_name}。{content_text}"
                
                if '。' in full_text:
                    first_period = full_text.index('。') + 1
                    subtitle = full_text[:first_period]
                    body_text = full_text[first_period:]
                    
                    # 添加加粗的小标题run
                    title_run = new_para.add_run(subtitle)
                    title_run.bold = True
                    # 完整设置字体（包括中文字体eastAsia）
                    title_run.font.size = Pt(16)
                    rPr = title_run._element.get_or_add_rPr()
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = rPr.makeelement(qn('w:rFonts'), {})
                        rPr.insert(0, rFonts)
                    rFonts.set(qn('w:hint'), 'eastAsia')
                    rFonts.set(qn('w:ascii'), '仿宋_GB2312')
                    rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
                    rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    rFonts.set(qn('w:cs'), '仿宋_GB2312')
                    
                    # 添加不加粗的正文run
                    body_run = new_para.add_run(body_text)
                    body_run.bold = False
                    # 完整设置字体（包括中文字体eastAsia）
                    body_run.font.size = Pt(16)
                    rPr = body_run._element.get_or_add_rPr()
                    rFonts = rPr.find(qn('w:rFonts'))
                    if rFonts is None:
                        rFonts = rPr.makeelement(qn('w:rFonts'), {})
                        rPr.insert(0, rFonts)
                    rFonts.set(qn('w:hint'), 'eastAsia')
                    rFonts.set(qn('w:ascii'), '仿宋_GB2312')
                    rFonts.set(qn('w:hAnsi'), '仿宋_GB2312')
                    rFonts.set(qn('w:eastAsia'), '仿宋_GB2312')
                    rFonts.set(qn('w:cs'), '仿宋_GB2312')
                    # 确保XML中明确设置bold为off
                    b_elem = rPr.find(qn('w:b'))
                    if b_elem is not None:
                        b_elem.set(qn('w:val'), '0')
                    else:
                        b_new = rPr.makeelement(qn('w:b'), {})
                        b_new.set(qn('w:val'), '0')
                        rPr.append(b_new)
                else:
                    run = new_para.add_run(full_text)
                    run.bold = True
                    run.font.size = Pt(16)
                    run.font.name = '仿宋_GB2312'
                
                print(f"[Report] 添加新段落: {idx + 1}.{dim_name}")
                last_para = new_para  # 更新last_para以便后续插入
    
    print(f"[Report] AI建议替换完成")


def generate_report(data_json, funnel_base64=None, ai_suggestions=None):
    """生成Word报告"""
    
    # 加载模板
    doc = Document(TEMPLATE_PATH)
    
    # 解析数据
    parsed = json.loads(data_json)
    # 修复：前端发送的格式是 {data: afReportData, funnelImage: ..., aiSuggestions: ...}
    # 所以需要从 parsed['data'] 获取实际报告数据
    data = parsed.get('data', parsed)
    print(f'[DEBUG] 后端解析: resultChurnRate={data.get("resultChurnRate")}, identityChurnRate={data.get("identityChurnRate")}')
    print(f'[DEBUG] 后端解析: transferRate={data.get("transferRate")}, completeRate={data.get("completeRate")}')
    
    # 替换占位符
    replace_all_placeholders(doc, data)
    
    # 替换AI建议（第三章）
    replace_ai_suggestions(doc, ai_suggestions)
    
    # 设置标题级别（一级目录）
    set_heading_outline_level(doc)
    
    # 添加漏斗图
    if funnel_base64:
        add_funnel_image(doc, funnel_base64)
    
    # 保存到字节流
    output = BytesIO()
    doc.save(output)
    output.seek(0)
    
    return output
